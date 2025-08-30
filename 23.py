import pandas as pd
from docx import Document

def create_labels(csv_file):
    try:
        df = pd.read_csv(csv_file)
        for index, row in df.iterrows():
            doc = Document()
            doc.add_paragraph(f"Name: {row['Name']}")
            doc.add_paragraph(f"Father's Name: {row['FatherName']}")
            doc.add_paragraph(f"Address: {row['Address']}")
            doc.add_paragraph(f"District: {row['District']}")
            doc.add_paragraph(f"State: {row['State']}")
            doc.add_paragraph(f"Pincode: {row['Pincode']}")
            doc.save(f"label_{index + 1}.docx")
        print("Labels generated successfully.")
    except Exception as e:
        print(f"Error: {e}")


create_labels("students.csv")
