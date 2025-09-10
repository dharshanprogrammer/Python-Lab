from docx import Document


class Student:
    def __init__(self, reg_no):
        self.reg_no = reg_no
        self.info_file = f"{reg_no}_info.txt"
        self.marks_file = f"{reg_no}_marks.txt"

    def save_personal_info(self, name, age, course):
        with open(self.info_file, 'w') as file:
            file.write(f"Name: {name}\n")
            file.write(f"Age: {age}\n")
            file.write(f"Course: {course}\n")
        print(f"Personal info saved to {self.info_file}")

    def save_marks(self, subject_marks):
        with open(self.marks_file, 'w') as file:
            for subject, marks in subject_marks.items():
                file.write(f"{subject}: {marks}\n")
        print(f"Marks saved to {self.marks_file}")

    def generate_report(self):
        document = Document()
        document.add_heading('Student Mark Statement', 0)

        document.add_heading('Personal Information:', level=1)
        with open(self.info_file, 'r') as file:
            document.add_paragraph(file.read())

        document.add_heading('Marks:', level=1)
        with open(self.marks_file, 'r') as file:
            document.add_paragraph(file.read())

        report_filename = f"{self.reg_no}_report.docx"
        document.save(report_filename)
        print(f"Report generated: {report_filename}")


# Example usage
if __name__ == "__main__":
    reg_no = input("Enter Student Registration Number: ")
    student = Student(reg_no)

    name = input("Enter Student Name: ")
    age = input("Enter Student Age: ")
    course = input("Enter Student Course: ")
    student.save_personal_info(name, age, course)

    subject_marks = {}
    n = int(input("Enter number of subjects: "))
    for _ in range(n):
        subject = input("Subject Name: ")
        marks = input("Marks: ")
        subject_marks[subject] = marks

    student.save_marks(subject_marks)
    student.generate_report()
