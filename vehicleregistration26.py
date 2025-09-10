from docx import Document

class Vehicle:
    def __init__(self, reg_no):
        self.reg_no = reg_no
        self.file_name = f"{reg_no}_vehicle.txt"

    def register_vehicle(self, owner_name, vehicle_type, model):
        with open(self.file_name, 'w') as file:
            file.write(f"Owner Name: {owner_name}\n")
            file.write(f"Vehicle Type: {vehicle_type}\n")
            file.write(f"Model: {model}\n")
        print(f"Vehicle details saved to {self.file_name}")

    def generate_rc_book(self):
        document = Document()
        document.add_heading('Digital RC Book', 0)
        with open(self.file_name, 'r') as file:
            document.add_paragraph(file.read())

        rc_filename = f"{self.reg_no}_rcbook.docx"
        document.save(rc_filename)
        print(f"RC Book generated: {rc_filename}")


# Example usage
if __name__ == "__main__":
    reg_no = input("Enter Vehicle Registration Number: ")
    vehicle = Vehicle(reg_no)

    owner_name = input("Enter Owner Name: ")
    vehicle_type = input("Enter Vehicle Type (Car/Truck/Bike): ")
    model = input("Enter Vehicle Model: ")

    vehicle.register_vehicle(owner_name, vehicle_type, model)
    vehicle.generate_rc_book()
